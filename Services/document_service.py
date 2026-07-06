from docx import Document

from Models.response import PlannerResponse


class DocumentService:

    def create_document(
        self,
        planner_response: PlannerResponse,
        filename: str
    ):

        document = Document()

        document.add_heading(
            "AI Generated Todo List",
            level=1
        )

        for index, task in enumerate(planner_response.tasks, start=1):

            document.add_heading(
                f"Task {index}: {task.title}",
                level=2
            )

            document.add_paragraph(
                f"Description: {task.description}"
            )

            document.add_paragraph(
                f"Duration: {task.duration}"
            )

            document.add_paragraph(
                f"Priority: {task.priority}"
            )

        document.save(filename)